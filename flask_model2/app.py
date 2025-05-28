from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import logging
import random
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import re
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Initialize Flask app
app = Flask(__name__)

# Enable CORS with environment variable
ALLOWED_ORIGINS = os.getenv('ALLOWED_ORIGINS', 'http://localhost:3000')
CORS(app, resources={r"/*": {"origins": ALLOWED_ORIGINS}})

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration for file handling
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PREPROCESSED_FOLDER = os.path.join(BASE_DIR, 'preprocessed')
PROCESSED_FOLDER = os.path.join(BASE_DIR, 'processed')

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PREPROCESSED_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PREPROCESSED_FOLDER'] = PREPROCESSED_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB file size limit

# Legal terms and responses
LEGAL_TERMS = [
    # Court and judicial terms
    "court", "judgment", "judge", "bench", "hon'ble", "honorable", "justice",
    "chief justice", "justice", "j.", "cj", "division bench", "full bench",
    "single bench", "coram", "jurisdiction", "original jurisdiction",
    "appellate jurisdiction", "writ jurisdiction", "revision", "review",
    "curative petition", "special leave petition", "slp", "civil appeal",
    "criminal appeal", "letters patent appeal", "lpa", "review petition",

    # Judgment document structure
    "case no.", "in the matter of", "versus", "vs", "v.", "petitioner",
    "respondent", "appellant", "appellee", "complainant", "accused",
    "defendant", "plaintiff", "applicant", "opposite party", "op",
    "intervenor", "amicus curiae", "next friend", "pro forma respondent",

    # Legal document sections
    "headnote", "citation", "facts", "issues", "arguments", "submissions",
    "contentions", "pleadings", "evidence", "exhibits", "affidavit",
    "deposition", "testimony", "witness", "examination", "cross-examination",
    "reexamination", "documents", "annexures", "schedules", "appendices",
    "preamble", "recitals", "operative portion", "ratio decidendi",
    "obiter dicta", "holding", "findings", "conclusions", "decision",
    "order", "decree", "final order", "interim order", "injunction",
    "stay", "bail", "remand", "custody", "parole", "probation",

    # Legal procedures
    "filing", "institution", "commencement", "cause of action",
    "limitation", "prescription", "res judicata", "lis pendens",
    "sub judice", "stare decisis", "precedent", "binding precedent",
    "persuasive precedent", "distinguished", "overruled", "reversed",
    "affirmed", "modified", "remanded", "disposed", "dismissed",
    "allowed", "partly allowed", "quashed", "set aside", "annulled",
    "struck down", "upheld", "sustained", "vacated", "withdrawn",

    # Constitutional law
    "constitution", "constitutional", "unconstitutional", "ultra vires",
    "intra vires", "basic structure", "fundamental rights", "directive principles",
    "fundamental duties", "writ", "habeas corpus", "mandamus", "prohibition",
    "certiorari", "quo warranto", "article 14", "article 19", "article 21",
    "article 32", "article 226", "article 136", "article 142", "article 144",
    "separation of powers", "judicial review", "rule of law", "due process",
    "equal protection", "reasonable restriction", "public interest",
    "doctrine of eclipse", "doctrine of severability", "colourable legislation",

    # Civil law
    "civil procedure code", "cpc", "order", "rule", "section", "appeal",
    "revision", "review", "execution", "decree", "judgment", "plaint",
    "written statement", "counter claim", "set off", "interlocutory",
    "interim relief", "injunction", "temporary injunction", "permanent injunction",
    "specific performance", "declaration", "damages", "compensation",
    "mesne profits", "restitution", "receiver", "commission", "discovery",
    "interrogatories", "admission", "denial", "affidavit", "ex parte",
    "ex parte decree", "ex parte order", "setting aside ex parte",

    # Criminal law
    "criminal procedure code", "crpc", "indian penal code", "ipc",
    "section 302", "section 304", "section 307", "section 376", "section 420",
    "section 498a", "bailable", "non-bailable", "cognizable", "non-cognizable",
    "fir", "charge sheet", "charges", "framing of charges", "discharge",
    "acquittal", "conviction", "sentence", "death sentence", "life imprisonment",
    "fine", "compensation", "probation", "parole", "remission", "commutation",
    "suspension", "anticipatory bail", "regular bail", "default bail",
    "custodial interrogation", "police custody", "judicial custody",
    "remand", "discharge", "compounding", "quashing", "stay", "suspension",

    # Evidence law
    "indian evidence act", "evidence", "proof", "burden of proof",
    "onus of proof", "standard of proof", "presumption", "rebuttable",
    "irrebuttable", "documentary evidence", "oral evidence", "expert evidence",
    "circumstantial evidence", "direct evidence", "hearsay", "confession",
    "admission", "dying declaration", "hostile witness", "leading question",
    "cross examination", "reexamination", "affidavit", "exhibit", "marking",
    "identification", "proof of document", "secondary evidence", "primary evidence",

    # Contract law
    "indian contract act", "contract", "agreement", "offer", "acceptance",
    "consideration", "competent parties", "free consent", "coercion",
    "undue influence", "fraud", "misrepresentation", "mistake", "void",
    "voidable", "unenforceable", "quasi contract", "specific performance",
    "damages", "compensation", "liquidated damages", "penalty", "rescission",
    "rectification", "restitution", "quantum meruit", "breach", "anticipatory",
    "actual breach", "remedies", "injunction", "declaration",

    # Property law
    "transfer of property act", "topa", "sale", "mortgage", "lease",
    "gift", "exchange", "license", "easement", "adverse possession",
    "prescription", "title", "ownership", "possession", "constructive possession",
    "joint possession", "co-ownership", "coparcenary", "partition", "will",
    "testament", "codicil", "probate", "letters of administration",
    "succession", "inheritance", "heir", "legatee", "devisee", "bequest",
    "gift", "settlement", "trust", "beneficiary", "trustee", "endowment",

    # Company law
    "companies act", "memorandum", "articles", "incorporation", "registration",
    "director", "managing director", "whole-time director", "independent director",
    "nominee director", "board of directors", "general meeting", "agm", "egm",
    "resolution", "ordinary resolution", "special resolution", "shareholder",
    "member", "share", "equity share", "preference share", "debenture",
    "charge", "mortgage", "lien", "floating charge", "fixed charge", "winding up",
    "voluntary winding up", "compulsory winding up", "liquidation", "official liquidator",
    "insolvency", "bankruptcy", "resolution professional", "liquidator",
    "winding up petition", "oppression and mismanagement", "nclt", "nclat",

    # Intellectual property
    "patent", "copyright", "trademark", "design", "geographical indication",
    "infringement", "passing off", "counterfeiting", "piracy", "plagiarism",
    "assignment", "license", "compulsory license", "royalty", "damages",
    "injunction", "anticipatory injunction", "permanent injunction",
    "account of profits", "seizure", "destruction", "rectification",

    # Labor and industrial law
    "industrial disputes act", "ida", "workman", "employer", "employee",
    "industrial dispute", "strike", "lockout", "layoff", "retrenchment",
    "closure", "transfer", "closure", "compensation", "gratuity", "bonus",
    "provident fund", "esic", "epf", "minimum wages", "equal remuneration",
    "sexual harassment", "disciplinary proceedings", "domestic enquiry",
    "punishment", "dismissal", "termination", "reinstatement", "back wages",
    "conciliation", "arbitration", "adjudication", "labour court",
    "industrial tribunal", "national tribunal", "collective bargaining",
    "settlement", "award", "implementation", "enforcement",

    # Tax law
    "income tax act", "gst", "vat", "customs", "excise", "service tax",
    "assessment", "reassessment", "scrutiny", "regular assessment",
    "best judgment assessment", "appeal", "revision", "rectification",
    "advance ruling", "settlement commission", "tax evasion", "tax avoidance",
    "penalty", "prosecution", "recovery", "attachment", "garnishee",
    "stay", "refund", "tribunal", "high court", "supreme court",

    # Arbitration
    "arbitration", "conciliation", "mediation", "award", "enforcement",
    "setting aside", "arbitrator", "umpire", "arbitral tribunal",
    "arbitration agreement", "seat", "venue", "jurisdiction", "competence",
    "competence-competence", "interim measures", "emergency arbitrator",
    "final award", "partial award", "interest", "costs", "challenge",
    "neutrality", "impartiality", "independence", "disclosure",

    # International law
    "treaty", "convention", "protocol", "united nations", "general assembly",
    "security council", "international court of justice", "pcij", "icj",
    "arbitration", "mediation", "good offices", "diplomatic protection",
    "state responsibility", "immunity", "sovereign immunity", "jurisdiction",
    "extradition", "mutual legal assistance", "human rights", "refugee",
    "asylum", "extra-territorial", "exhaustion of local remedies",

    # Legal maxims
    "actus reus", "mens rea", "audi alteram partem", "nemo judex in causa sua",
    "res ipsa loquitur", "uberrima fides", "caveat emptor", "stare decisis",
    "obiter dicta", "ratio decidendi", "ignorantia juris non excusat",
    "de minimis non curat lex", "expressio unius est exclusio alterius",
    "ejusdem generis", "noscitur a sociis", "pari materia", "in pari delicto",
    "volenti non fit injuria", "damnum sine injuria", "injuria sine damno",
    "qui facit per alium facit per se", "respondeat superior", "volenti non fit injuria",

    # Case-specific terms (from your judgment)
    "appeal no. 790 of 1957", "civil misc writ no. 280 of 1950",
    "u.p. industrial disputes act", "xxviii of 1947", "court of inquiry",
    "allahabad high court", "mudholkar j.", "bhargava j.", "sapru j.",
    "state of uttar pradesh", "indian sugar millers association",
    "indian national sugar mills workers federation", "sugar factories",
    "bonus payment", "retaining allowance", "seasonal workmen", "clerical staff",
    "industrial dispute", "strike notice", "court of inquiry", "gazette notification",
    "writ petition", "article 226", "article 133", "full bench", "constitutional validity",
    "ultra vires", "discrimination", "arbitrary", "public interest", "emergency",
    "prospective", "retrospective", "minimum wages act", "collective bargaining",
    "terms of employment", "conditions of employment", "mandamus", "certificate",
    "special leave petition", "constitution bench",

    # Terms related to user questions and queries
    "explain", "clarify", "interpret", "define", "what does", "meaning of",
    "how to", "procedure for", "requirements for", "eligibility for",
    "criteria for", "difference between", "similarities between",
    "compare", "contrast", "examples of", "types of", "categories of",
    "applicability of", "scope of", "limitations of", "exceptions to",
    "validity of", "enforceability of", "consequences of", "penalty for",
    "remedy for", "solution for", "process for", "steps to", "guide to",
    "analysis of", "breakdown of", "summary of", "overview of",
    "key points", "main arguments", "legal basis", "grounds for",
    "justification for", "rationale behind", "purpose of", "intent behind",
    "objective of", "effect of", "impact of", "implications of",
    "significance of", "importance of", "relevance of", "connection between",
    "relationship between", "correlation between", "cause of", "effect of",
    "reason for", "basis for", "foundation of", "principle behind",
    "doctrine of", "theory of", "concept of", "aspects of", "elements of",
    "components of", "factors in", "considerations for", "requirements of",
    "conditions for", "terms of", "provisions of", "clauses in",
    "sections in", "articles in", "rules in", "regulations in",
    "guidelines for", "standards for", "benchmarks for", "precedents for",
    "case law on", "jurisprudence on", "legal opinion on", "view on",
    "position on", "stance on", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "application of",
    "implementation of", "execution of", "enforcement of", "compliance with",
    "adherence to", "obligation to", "right to", "entitlement to",
    "privilege of", "immunity from", "exception to", "exemption from",
    "derogation from", "deviation from", "variation of", "modification of",
    "amendment to", "revision of", "update to", "change in",
    "development in", "trend in", "pattern in", "practice of",
    "custom of", "usage of", "tradition of", "convention of",
    "norm of", "standard of", "measure of", "test for",
    "criteria for", "benchmark for", "yardstick for", "indicator of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "approval of",
    "authorization of", "sanction of", "ratification of", "endorsement of",
    "support for", "opposition to", "objection to", "challenge to",
    "appeal against", "review of", "reconsideration of", "revision of",
    "reformation of", "rectification of", "correction of", "amendment of",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "paraphrase of", "summary of", "abstract of", "synopsis of",
    "outline of", "overview of", "introduction to", "background of",
    "context of", "framework of", "structure of", "organization of",
    "hierarchy of", "classification of", "categorization of", "typology of",
    "taxonomy of", "nomenclature of", "terminology of", "vocabulary of",
    "glossary of", "dictionary of", "lexicon of", "thesaurus of",
    "encyclopedia of", "compendium of", "digest of", "manual of",
    "handbook of", "guidebook of", "textbook of", "treatise on",
    "monograph on", "dissertation on", "thesis on", "paper on",
    "article on", "essay on", "commentary on", "annotation of",
    "exegesis of", "hermeneutics of", "interpretation of", "construction of",
    "reading of", "analysis of", "examination of", "investigation of",
    "inquiry into", "research on", "study of", "survey of",
    "report on", "finding of", "conclusion of", "recommendation of",
    "suggestion for", "proposal for", "plan for", "strategy for",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for","industrial","act","payment","law","court","bonus","order","government"
]
IRRELEVANT_RESPONSES = [
    "This question appears unrelated to the legal judgment document.",
    "The system only answers questions specifically about the uploaded court judgment.",
    "Your question doesn't appear relevant to this legal document.",
    "For questions about this specific judgment, please reference the case details.",
    "I can only answer questions about the legal judgment document."
]

# Lazy-load embeddings model
embeddings = None
vectorstore_cache = {}

def load_embeddings():
    global embeddings
    if embeddings is None:
        logger.info("⏳ Loading embeddings model...")
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={
                'normalize_embeddings': True,
                'batch_size': 4
            }
        )
        logger.info("✅ Embeddings model loaded successfully!")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def preprocess_text(text):
    try:
        if not text or not isinstance(text, str):
            return ""
            
        text = re.sub(r'Page\s*\d+\s*of\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\d+\n', '\n', text)
        
        noise_patterns = [
            r'This\s+document\s+was\s+signed\s+electronically.*',
            r'Electronic\s+signature.*',
            r'IN\s+THE\s+(?:SUPREME\s+)?COURT\s+OF\s+.*\n',
            r'CASE\s+NO[.:].*\n',
            r'BEFORE[.:].*\n',
            r'JUDGMENT\s+RESERVED\s+ON[.:].*\n',
            r'PRESENT[.:].*\n'
        ]
        
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if (re.fullmatch(r'\d+', line) or 
                len(line) < 25 or 
                line.startswith(('Sl', '©', 'Page', 'http'))):
                continue
                
            cleaned_lines.append(line)
        
        return '\n\n'.join(cleaned_lines) or ""
    except Exception as e:
        logger.error(f"Error in preprocessing: {str(e)}")
        return text if isinstance(text, str) else ""

def get_processed_filename(original_filename):
    base = os.path.splitext(os.path.basename(original_filename))[0]
    return f"processed_{secure_filename(base)}.txt"

def get_processed_path(filename):
    processed_filename = get_processed_filename(filename)
    return os.path.normpath(os.path.join(app.config['PROCESSED_FOLDER'], processed_filename))

def extract_text_from_file(filepath, filename):
    try:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            return text.strip()
        
        elif ext in ['.doc', '.docx']:
            doc = docx.Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs if para.text)
        
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        
        raise ValueError(f"Unsupported file type: {ext}")
    
    except Exception as e:
        logger.error(f"Extraction failed for {filename}: {str(e)}")
        raise

def create_vector_store(text, filename):
    try:
        # Safeguard: Check if embeddings is loaded
        if embeddings is None:
            raise ValueError("Embeddings model is not loaded. Call load_embeddings() first.")
        
        processed_path = get_processed_path(filename)
        os.makedirs(os.path.dirname(processed_path), exist_ok=True)
        
        with open(processed_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        if not os.path.exists(processed_path):
            raise IOError(f"Failed to create processed file at {processed_path}")
        
        with open(processed_path, 'r', encoding='utf-8') as f:
            text = f.read()
        documents = [Document(page_content=text, metadata={"source": processed_path})]
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(documents)
        
        vectorstore = FAISS.from_documents(chunks, embeddings)
        return vectorstore, None
    except Exception as e:
        logger.error(f"Error in create_vector_store: {str(e)}")
        return None, str(e)

def contains_legal_terms(text):
    if not text:
        return False
    text_lower = text.lower()
    return sum(term in text_lower for term in LEGAL_TERMS) >= 2

def is_relevant_response(query, docs_and_scores):
    if not docs_and_scores or not any(score >= 0.8 for _, score in docs_and_scores):
        return False

    query_lower = query.lower()
    if not (query_lower.endswith('?') or 
            any(word in query_lower for word in ['who', 'what', 'when', 'where', 'why', 'how', 'explain'])):
        return False

    for doc, _ in docs_and_scores:
        if contains_legal_terms(doc.page_content) and contains_legal_terms(query):
            return True
    return False

def format_answer(doc, score):
    content = doc.page_content
    return {
        "content": content,
        "score": float(score)
    }

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "embeddings_loaded": embeddings is not None,
        "device": "cpu"
    })

@app.route('/upload', methods=['POST'])
def upload_document():
    if embeddings is None:
        try:
            load_embeddings()
        except Exception as e:
            return jsonify({"error": "Embeddings model failed to load", "status": "error"}), 503
    
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded", "status": "error"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file", "status": "error"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({
            "error": "Invalid file type. Allowed: pdf, doc, docx, txt",
            "status": "error"
        }), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(filepath)
        
        raw_text = extract_text_from_file(filepath, filename)
        if not raw_text:
            return jsonify({
                "error": "Could not extract text from document",
                "status": "error"
            }), 400
            
        cleaned_text = preprocess_text(raw_text)
        vectorstore, error = create_vector_store(cleaned_text, filename)
        if error:
            return jsonify({"error": error, "status": "error"}), 500
        
        vectorstore_cache[filename] = vectorstore
        
        return jsonify({
            "message": "Document processed successfully",
            "filename": filename,
            "status": "success"
        })
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return jsonify({
            "error": "Document processing failed",
            "details": str(e),
            "status": "error"
        }), 500
    finally:
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception as e:
                logger.error(f"Error removing file {filepath}: {str(e)}")

@app.route('/ask', methods=['POST'])
def ask_question():
    try:
        # Ensure embeddings are loaded before proceeding
        if embeddings is None:
            try:
                load_embeddings()
            except Exception as e:
                return jsonify({"error": "Embeddings model failed to load", "status": "error"}), 503

        if 'file' not in request.files:
            return jsonify({
                "error": "Missing file",
                "status": "error"
            }), 400

        file = request.files['file']
        question = request.form.get('question', '').strip()

        if not file or file.filename == '':
            return jsonify({"error": "No file selected", "status": "error"}), 400

        if len(question.split()) < 3:
            return jsonify({
                "error": "Please ask a more detailed question (minimum 3 words)",
                "status": "success"
            }), 200

        filename = secure_filename(file.filename)
        vectorstore = vectorstore_cache.get(filename)
        
        if vectorstore is None:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            raw_text = extract_text_from_file(filepath, filename)
            if not raw_text:
                return jsonify({
                    "error": "Could not extract text from document",
                    "status": "error"
                }), 400

            cleaned_text = preprocess_text(raw_text)
            vectorstore, error = create_vector_store(cleaned_text, filename)
            if error:
                return jsonify({"error": error, "status": "error"}), 500
            
            vectorstore_cache[filename] = vectorstore  # Cache the vectorstore
            
            if os.path.exists(filepath):
                os.remove(filepath)

        docs_and_scores = vectorstore.similarity_search_with_score(question, k=5)
        
        if not is_relevant_response(question, docs_and_scores):
            return jsonify({
                "answer": random.choice(IRRELEVANT_RESPONSES),
                "sections": [],
                "isRelevant": False,
                "filename": filename,
                "status": "success"
            })
        
        relevant_sections = []
        for doc, score in docs_and_scores:
            if score >= 0.8 and contains_legal_terms(doc.page_content):
                relevant_sections.append(format_answer(doc, score))

        return jsonify({
            "answer": "Here are the relevant sections from the document:",
            "sections": relevant_sections,
            "filename": filename,
            "isRelevant": True,
            "status": "success"
        })
    except Exception as e:
        logger.error(f"Error answering question: {str(e)}")
        return jsonify({
            "error": "Failed to process question",
            "details": str(e),
            "status": "error"
        }), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8001))
    app.run(host='0.0.0.0', port=port)